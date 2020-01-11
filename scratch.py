# coding=utf-8
import xlrd
import random
import docx
from docx.oxml.ns import qn
from docx.shared import Cm,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
def tkdaxhx(str1):
    num = 0
    sum = 0
    wz = []
    for x in str1:
        if (x == '_' or x == '*'):

            if sum % 2 == 0:
                wz.append([num])
            else:
                wz[int(sum / 2)].append(num)
            sum = sum + 1
        num = num + 1
    # print(wz)
    gb=0
    timu=doc.add_paragraph()
    timu.paragraph_format.space_after = Pt(1)
    timu.paragraph_format.space_before = Pt(1)

    str1=str1.replace('_',' ')
    str1 = str1.replace('*', ' ')
    for y in wz:
        timuRun=timu.add_run(str1[gb:y[0]])
        timuRun.font.size = (Pt(10.5))
        try:
            # print(str[y[0]:y[1]+1])
            timuRun = timu.add_run(str1[y[0]:y[1] + 1])
            timuRun.font.size = (Pt(10.5))
            timuRun.font.underline=True
            gb=y[1]+1
        except:
            #print("生成填空题第" + str(str1[0]) + "题时发现单括号或单下划线，请手动调整")
            gb=y[0]
    if gb<len(str1):
        timuRun = timu.add_run(str1[gb:])
        timuRun.font.size = (Pt(10.5))
def tkdacl(str1):
    num=0
    sum=0
    wz=[]
    for x in str1:
        if (x=='_' or x=='*'):

            if sum%2==0:
                wz.append([num])
            else:
                wz[int(sum/2)].append(num)
            sum = sum + 1
        num=num+1
    #print(wz)
    for y in wz:
        kgs=0
        tihuan='_'
        try:
            while kgs<(y[1]-y[0]):
                tihuan=tihuan+'_'
                kgs=kgs+1
            #print(str1[y[0]:y[1]+1])
            str1=str1.replace(str1[y[0]:y[1]+1],tihuan)


        except:
            print("生成填空题第"+str(str1.split('.')[0])+"题时发现单括号或单下划线，请手动调整")
    str1=str1.replace("_","__")
    return str1

def findNum(str,fg):
    return len(str.split(fg))
def xccfxhx(str):
    while str.find('__')!=-1:
        str=str.replace('__','_')
    return str
def kkhcl(str):
    num=0
    sy=-1

    if str.find('（')!=-1:
        sy=str.find('（')
        while num<len(str)-1:
            if str[num]=='（':
                kgjs=0
                sy=str.find('（')
                num = num + 1
                while num<len(str) and (((str[num]==' ' )or (str[num]=='）' ))) :
                    if str[num]=='）':
                        str=str[:sy] + '_' + str[num+1:]
                        num=num-kgjs+1
                    else:
                        kgjs=kgjs+1
                    num = num + 1
            num = num + 1



    return str





def read_xlrd(excelFile):
    data = xlrd.open_workbook(excelFile)
    table = data.sheet_by_index(0)
    for rowNum in range(table.nrows):
        rowVale = table.row_values(rowNum)
        for colNum in range(table.ncols):
            if rowNum > 0 and colNum == 0:
                tk.append([])
                # print(int(rowVale[0]))
                tk[rowNum].append(int(rowVale[0]))
            else:
                # print(rowVale[colNum])
                tk[rowNum].append(rowVale[colNum])
                pass
            # print("---------------")


def chuti(tixing, shuliang,fenshu):
    printCon=''
    printCon1 = ''
    if tixing == 1:

        tx1 = []
        num = 0
        for timu in tk:
            if timu[1] == '单项选择':
                tx1.append(timu)
        # print(tx1)
        ct = random.sample(tx1, shuliang)
        #print(ct)
        ctth = 0
        for cttm in ct:
            ctth = ctth + 1
            cttm[2].lstrip()
            cttm[2].rstrip()
            #print(cttm[2].find('（ *）'))
            #print(cttm[2].split('（ *）'))
            #print(cttm[2])
            while cttm[2][len(cttm[2])-1]=='\n':
                cttm[2]=cttm[2][:-1]
            if cttm[2].find('( ')!=-1:
                #print('aaa')
                cttm2 = cttm[2].replace('( ', '(  ' + '  ', 1)
                cttm[2]=cttm[2].replace('( ','(  '+cttm[7] +'  ',1)

            elif cttm[2].find('()')!=-1:
                #print('bbb')
                cttm2 = cttm[2].replace('()', '（  ' + '  ', 1)
                cttm[2] = cttm[2].replace('()', '（  ' + cttm[7] + '  ',1)

            elif cttm[2].find('（）')!=-1:
                #print('ccc')
                cttm2 = cttm[2].replace('（）', '（  ' + '  ）', 1)
                cttm[2] = cttm[2].replace('（）', '（  ' + cttm[7] + '  ）',1)

            elif cttm[2].find('（ ') !=-1:
                #print('ddd')
                cttm2 = cttm[2].replace('（ ', '（  ' +  '  ', 1)
                cttm[2] = cttm[2].replace('（ ', '（  ' + cttm[7] + '  ',1)
            elif cttm[2].find('_') != -1:
                cttm[2] = cttm[2].replace('__', '_')
                cttm[2] = cttm[2].replace('__', '_')
                cttm[2] = cttm[2].replace('__', '_')
                cttm[2] = cttm[2].replace('__', '_')
                cttm[2] = cttm[2].replace('__', '_')
                #print('eee')
                cttm2 = cttm[2].replace('_', '（  ' +  '  ）', 1)
                cttm[2] = cttm[2].replace('_', '（  ' + cttm[7] + '  ）',1)
            elif cttm[2].find('  ') != -1:
                #print('fff')
                cttm2 = cttm[2].replace('  ', '（  '  + '  ）', 1)
                cttm[2] = cttm[2].replace('  ', '（  ' + cttm[7] + '  ）', 1)
            else:
                #print('ggg')
                cttm2 = cttm[2] + '（  '  + '  ）'
                cttm[2] = cttm[2]+'（  '+cttm[7]+'  ）'
            printCon=printCon+str(ctth) + '.' + cttm[2] + '   ' +  '\n' + 'A、' + str(cttm[3]) + '  B、' + str(
                cttm[4]) + '  C、' + str(cttm[5]) + '  D、' + str(cttm[6])+'\n'
            printCon1=printCon1+str(ctth) + '.' + cttm2 + '   ' +  '\n' + 'A、' + str(cttm[3]) + '  B、' + str(
                cttm[4]) + '  C、' + str(cttm[5]) + '  D、' + str(cttm[6])+'\n'
    elif tixing == 2:
        tx2 = []
        num = 0
        for timu in tk:
            if timu[1] == '填空题':
                tx2.append(timu)
        # print(tx1)
        ct = random.sample(tx2, shuliang)
        #print(ct)
        ctth = 0
        for cttm in ct:
            # cttm[2] = ' '.join(cttm[2].split())
            ctth = ctth + 1
            cttm[2].lstrip()
            while cttm[2][len(cttm[2])-1]=='\n':
                cttm[2]=cttm[2][:-1]
            #cttm[2].rstrip()
            #print(cttm[2])
            #cttm[2] = cttm[2].replace('(', '_')
            #cttm[2] = cttm[2].replace(')', '_')
            #cttm[2] = cttm[2].replace('（', '_')
            #cttm[2] = cttm[2].replace('）', '_')
            #cttm[2] = cttm[2].replace('( ', '_')
            #cttm[2] = cttm[2].replace(' )', '_')
            #cttm[2] = cttm[2].replace('（ ', '_')
            #cttm[2] = cttm[2].replace(' ）', '_')
            # '_'.join(cttm[2].split('_'))

            if kkhcl(cttm[2])==-1:
                pass
            else:
                cttm[2]=kkhcl(cttm[2])
                #cttm[2][kkh(cttm[2])[0]]='_'
                #cttm[2][kkh(cttm[2])[1]] = '_'
                #cttm[2]=cttm[2][:kkh(cttm[2])[0]]+'_'+cttm[2][kkh(cttm[2])[1]:]
            cttm[2] = cttm[2].replace('  ', '_')
            cttm[2] = xccfxhx(cttm[2])
            cttm[7]=str(cttm[7])
            #print(cttm[2])
            if len(cttm[2].split('_'))==1:
                printCon = str(ctth) + '.' + cttm[2] + '  ' + '_' + str(cttm[7]) + '_'
                tkdaxhx(printCon)
                printCon1 = printCon1 + tkdacl(str(ctth) + '.' + cttm[2] + '  ' + '_' + str(cttm[7]) + '_' + '\n')
            elif len(cttm[2].split('_'))==2:
                cttm[2] = cttm[2].replace('_', '_' + str(cttm[7]) + '_', 1)
                printCon=str(ctth) + '.' + cttm[2]
                tkdaxhx(printCon)
                printCon1=printCon1+tkdacl(str(ctth) + '.' + cttm[2]+'\n')
            elif len(cttm[2].split('_'))==3:
                #print(cttm)

                if cttm[7].find('；')!=-1:
                    anw=cttm[7].split("；",1)
                    cttm[2] = cttm[2].replace('_', '*' + str(anw[0]) + '*', 1)
                    cttm[2] = cttm[2].replace('_', '_' + str(anw[1]) + '_', 1)

                    printCon=str(ctth) + '.' + cttm[2]
                    tkdaxhx(printCon)
                    printCon1=printCon1+tkdacl(str(ctth) + '.' + cttm[2]+'\n')
                elif cttm[7].find('，')!=-1:
                    anw=cttm[7].split("，",1)
                    cttm[2] = cttm[2].replace('_', '*' + str(anw[0]) + '*', 1)
                    cttm[2] = cttm[2].replace('_', '_' + str(anw[1]) + '_', 1)
                    printCon=str(ctth) + '.' + cttm[2]
                    tkdaxhx(printCon)
                    printCon1=printCon1+tkdacl(str(ctth) + '.' + cttm[2]+'\n')
                elif cttm[7].find('、') != -1:
                    anw = cttm[7].split("、", 1)
                    cttm[2] = cttm[2].replace('_', '*' + str(anw[0]) + '*', 1)
                    cttm[2] = cttm[2].replace('_', '_' + str(anw[1]) + '_', 1)
                    printCon=str(ctth) + '.' + cttm[2]
                    tkdaxhx(printCon)
                    printCon1=printCon1+tkdacl(str(ctth) + '.' + cttm[2]+'\n')
                elif cttm[7].find(',') != -1:
                    anw = cttm[7].split(",", 1)
                    cttm[2] = cttm[2].replace('_', '*' + str(anw[0]) + '*', 1)
                    cttm[2] = cttm[2].replace('_', '_' + str(anw[1]) + '_', 1)
                    printCon=printCon+str(ctth) + '.' + cttm[2]
                    tkdaxhx(printCon)
                    printCon1=printCon1+tkdacl(str(ctth) + '.' + cttm[2]+'\n')
                else:
                    cttm[2] = cttm[2].replace('_', '_' + str(cttm[7]) + '_', 1)
                    printCon=str(ctth) + '.' + cttm[2]
                    tkdaxhx(printCon)
                    printCon1=printCon1+tkdacl(str(ctth) + '.' + cttm[2]+'\n')
            elif len(cttm[2].split('_')) == 4:
                if cttm[7].find('；')!=-1:
                    anw=cttm[7].split("；")
                    cttm[2] = cttm[2].replace('_', '*' + str(anw[0]) + '*', 1)
                    cttm[2] = cttm[2].replace('_', '*' + str(anw[1]) + '*', 1)
                    try:
                        cttm[2] = cttm[2].replace('_', '*' + str(anw[2]) + '*', 1)
                    except:
                        pass
                    #print(str(ctth) + '.' + cttm[2])
                if cttm[7].find('，')!=-1:
                    anw=cttm[7].split("，")
                    cttm[2] = cttm[2].replace('_', '*' + str(anw[0]) + '*', 1)
                    cttm[2] = cttm[2].replace('_', '*' + str(anw[1]) + '*', 1)
                    try:
                        cttm[2] = cttm[2].replace('_', '*' + str(anw[2]) + '*', 1)
                    except:
                        pass
                    #print(str(ctth) + '.' + cttm[2])
                if cttm[7].find('、') != -1:
                    anw = cttm[7].split("、")
                    cttm[2] = cttm[2].replace('_', '*' + str(anw[0]) + '*', 1)
                    cttm[2] = cttm[2].replace('_', '*' + str(anw[1]) + '*', 1)
                    try:
                        cttm[2] = cttm[2].replace('_', '*' + str(anw[2]) + '*', 1)
                    except:
                        pass
                    #print(str(ctth) + '.' + cttm[2])
                if cttm[7].find(',') != -1:
                    anw = cttm[7].split(",")
                    cttm[2] = cttm[2].replace('_', '*' + str(anw[0]) + '*', 1)
                    cttm[2] = cttm[2].replace('_', '*' + str(anw[1]) + '*', 1)
                    try:
                        cttm[2] = cttm[2].replace('_', '*' + str(anw[2]) + '*', 1)
                    except:
                        pass
                    #print(str(ctth) + '.' + cttm[2])

                cttm[2] = cttm[2].replace('_', '_' + str(cttm[7]) + '_', 1)
                printCon1=printCon1+tkdacl(str(ctth) + '.' + cttm[2]+'\n')
                printCon =str(ctth) + '.' + cttm[2]+'\n'
                tkdaxhx(printCon)
            #printCon = printCon.replace('*', '_')
        #printCon1=tkdacl(printCon)




    elif tixing == 3:

        tx3 = []
        num = 0
        for timu in tk:
            if timu[1] == '判断题':
                tx3.append(timu)
        #print(tx3)
        ct = random.sample(tx3, shuliang)
        #print(ct)
        ctth = 0
        for cttm in ct:
            ctth = ctth + 1
            while cttm[2][len(cttm[2])-1]=='\n':
                cttm[2]=cttm[2][:-1]
            if cttm[7] == 'Y':
                printCon=printCon+str(ctth) + '.' + cttm[2]+'   ' + '(是)\n'
                printCon1 = printCon1 + str(ctth) + '.' + cttm[2] + '   ' + '(   )\n'
            if cttm[7] == 'N':
                printCon=printCon+str(ctth) + '.' + cttm[2] + '   ' + '(否)\n'
                printCon1 = printCon1 + str(ctth) + '.' + cttm[2] + '   ' + '(   )\n'

    elif tixing == 4:

        tx4 = []
        num = 0
        for timu in tk:
            if timu[1] == '简答题':
                tx4.append(timu)
        #print(tx3)
        ct = random.sample(tx4, shuliang)
        #print(ct)
        ctth = 0
        for cttm in ct:
            ctth = ctth + 1
            while cttm[2][len(cttm[2])-1]=='\n':
                cttm[2]=cttm[2][:-1]
            printCon=printCon+str(ctth) + '.' + cttm[2]+'\n' + '答：'+'\n'+cttm[7]+'\n\n'
            printCon1 = printCon1 + str(ctth) + '.' + cttm[2] + '\n' + '答：' + '\n\n\n\n\n\n\n'
    elif tixing == 5:

        tx5 = []
        num = 0
        for timu in tk:
            if timu[1] == '多选题':
                tx5.append(timu)
        # print(tx1)
        ct = random.sample(tx5, shuliang)
        #print(ct)
        ctth = 0
        for cttm in ct:
            ctth = ctth + 1
            cttm[2].lstrip()
            cttm[2].rstrip()
            #print(cttm[2].find('（ *）'))
            #print(cttm[2].split('（ *）'))
            #print(cttm[2])
            while cttm[2][len(cttm[2])-1]=='\n':
                cttm[2]=cttm[2][:-1]
            if cttm[2].find('( ')!=-1:
                #print('aaa')
                cttm2 = cttm[2].replace('( ', '(  ' + '  ', 1)
                cttm[2]=cttm[2].replace('( ','(  '+cttm[7] +'  ',1)

            elif cttm[2].find('()')!=-1:
                #print('bbb')
                cttm2 = cttm[2].replace('()', '（  ' + '  ', 1)
                cttm[2] = cttm[2].replace('()', '（  ' + cttm[7] + '  ',1)

            elif cttm[2].find('（）')!=-1:
                #print('ccc')
                cttm2 = cttm[2].replace('（）', '（  ' + '  ）', 1)
                cttm[2] = cttm[2].replace('（）', '（  ' + cttm[7] + '  ）',1)

            elif cttm[2].find('（ ') !=-1:
                #print('ddd')
                cttm2 = cttm[2].replace('（ ', '（  ' +  '  ', 1)
                cttm[2] = cttm[2].replace('（ ', '（  ' + cttm[7] + '  ',1)
            elif cttm[2].find('_') != -1:
                cttm[2] = cttm[2].replace('__', '_')
                cttm[2] = cttm[2].replace('__', '_')
                cttm[2] = cttm[2].replace('__', '_')
                cttm[2] = cttm[2].replace('__', '_')
                cttm[2] = cttm[2].replace('__', '_')
                #print('eee')
                cttm2 = cttm[2].replace('_', '（  ' +  '  ）', 1)
                cttm[2] = cttm[2].replace('_', '（  ' + cttm[7] + '  ）',1)
            elif cttm[2].find('  ') != -1:
                #print('fff')
                cttm2 = cttm[2].replace('  ', '（  '  + '  ）', 1)
                cttm[2] = cttm[2].replace('  ', '（  ' + cttm[7] + '  ）', 1)
            else:
                #print('ggg')
                cttm2 = cttm[2] + '（  '  + '  ）'
                cttm[2] = cttm[2]+'（  '+cttm[7]+'  ）'
            printCon=printCon+str(ctth) + '.' + cttm[2] + '   ' +  '\n' + 'A、' + str(cttm[3]) + '  B、' + str(
                cttm[4]) + '  C、' + str(cttm[5]) + '  D、' + str(cttm[6])+'\n'
            printCon1=printCon1+str(ctth) + '.' + cttm2 + '   ' +  '\n' + 'A、' + str(cttm[3]) + '  B、' + str(
                cttm[4]) + '  C、' + str(cttm[5]) + '  D、' + str(cttm[6])+'\n'
    if tixing!=2:
        timu=doc.add_paragraph()
        timuRun=timu.add_run(printCon)
        timuRun.font.size = (Pt(10.5))
    timu1=doc1.add_paragraph()
    timuRun1=timu1.add_run(printCon1)
    timuRun1.font.size=(Pt(10.5))
    return printCon


if __name__ == '__main__':
    tk = [[]]
    daxie=['一','二','三','四','五','六','七','八','九']
    excelFile = '题库.xls'
    read_xlrd(excelFile=excelFile)
    doc=docx.Document()
    doc1=docx.Document()
    doc.styles['Normal'].font.name='宋体'
    doc.styles['Normal'].paragraph_format.space_after = Pt(5)
    doc.styles['Normal'].paragraph_format.space_before = Pt(5)
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    doc1.styles['Normal'].font.name='宋体'
    doc1.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    doc1.styles['Normal'].paragraph_format.space_after = Pt(5)
    doc1.styles['Normal'].paragraph_format.space_before = Pt(5)
    print("请输入试卷名称：")
    name=input()
    biaoti=doc.add_paragraph()
    biaotiRun=biaoti.add_run(name)
    biaotiRun.font.bold=True
    biaotiRun.font.size=Pt(18)
    biaoti_format=biaoti.paragraph_format
    biaoti_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fenshu=doc.add_table(rows=3,cols=0,style="Table Grid")
    fenshu.autofit = True
    fenshu.add_column(Cm(2)).cells[2].text = '得分'
    biaoti1 = doc1.add_paragraph()
    biaotiRun1 = biaoti1.add_run(name)
    biaotiRun1.font.bold = True
    biaotiRun1.font.size = Pt(18)
    biaoti_format1 = biaoti1.paragraph_format
    biaoti_format1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fenshu1 = doc1.add_table(rows=3, cols=0, style="Table Grid")
    fenshu1.autofit = True
    fenshu1.add_column(Cm(2)).cells[2].text = '得分'
    #fenshu.add_column(1).cells[0].text = '题号'
    #fenshu.cell(1,0).text='得分'
    #fenshu.cell(0, 1).text = '题号'
    th=[]
    kaoshengInfo=doc.add_paragraph()
    kaoshengInfoRun=kaoshengInfo.add_run('\n基地_________________姓名___________   ____年____月____日')
    kaoshengInfoRun.font.size = Pt(12)
    kaoshengInfo_format = kaoshengInfo.paragraph_format
    kaoshengInfo_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    zhuyishixiang = doc.add_paragraph()
    zhuyishixiangRun=zhuyishixiang.add_run('试卷说明：')
    zhuyishixiangRun.font.bold=True
    zhuyishixiangRun = zhuyishixiang.add_run('1.本试卷共计四大题。考试用时60分钟。\n          2.卷面分值100分，60分合格。\n          3.试卷用黑色或蓝色水笔、圆珠笔作答。简答题试卷正面如果书写不完在试卷反面作答。')
    zhuyishixiang_format = zhuyishixiang.paragraph_format
    zhuyishixiang_format.alignment=WD_ALIGN_PARAGRAPH.LEFT


    kaoshengInfo1=doc1.add_paragraph()
    kaoshengInfoRun1=kaoshengInfo1.add_run('\n基地_________________姓名___________   ____年____月____日')
    kaoshengInfoRun1.font.size = Pt(12)
    kaoshengInfo_format1 = kaoshengInfo1.paragraph_format
    kaoshengInfo_format1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    zhuyishixiang1 = doc1.add_paragraph()
    zhuyishixiangRun1=zhuyishixiang1.add_run('试卷说明：')
    zhuyishixiangRun1.font.bold=True
    zhuyishixiangRun1 = zhuyishixiang1.add_run('1.本试卷共计四大题。考试用时60分钟。\n          2.卷面分值100分，60分合格。\n          3.试卷用黑色或蓝色水笔、圆珠笔作答。简答题试卷正面如果书写不完在试卷反面作答。')
    zhuyishixiang_format1 = zhuyishixiang1.paragraph_format
    zhuyishixiang_format1.alignment=WD_ALIGN_PARAGRAPH.LEFT
    print('请输入填空题数量')
    thSum=0
    tkNum=int(input())
    if tkNum!=0:
        print('请输入填空题分数')
        tkFs=int(input())
        th.append(['填空',tkFs])
        timubiaoti=doc.add_paragraph()
        xiaofen=tkFs/tkNum
        timubiaotiRun=timubiaoti.add_run(daxie[thSum]+'、填空题'+'（每题'+str(xiaofen)+'分，共计'+str(tkFs)+'分）')
        timubiaotiRun.font.bold=True

        timubiaoti1=doc1.add_paragraph()
        timubiaotiRun1=timubiaoti1.add_run(daxie[thSum]+'、填空题'+'（每题'+str(xiaofen)+'分，共计'+str(tkFs)+'分）')
        timubiaotiRun1.font.bold=True
        thSum=thSum+1
        chuti(2, tkNum,tkFs)
    print('请输入单选题数量')
    dxNum=int(input())
    if dxNum!=0:
        print('请输入单选题分数')
        dxFs = int(input())
        th.append(['单选',dxFs])
        timubiaoti=doc.add_paragraph()
        xiaofen=dxFs/dxNum
        timubiaotiRun=timubiaoti.add_run(daxie[thSum]+'、单选题'+'（每题'+str(xiaofen)+'分，共计'+str(dxFs)+'分）')
        timubiaotiRun.font.bold=True

        timubiaoti1=doc1.add_paragraph()
        timubiaotiRun1=timubiaoti1.add_run(daxie[thSum]+'、单选题'+'（每题'+str(xiaofen)+'分，共计'+str(dxFs)+'分）')
        timubiaotiRun1.font.bold=True
        thSum = thSum + 1
        chuti(1,dxNum,dxFs)
    print('请输入多选题数量')
    duoxNum = int(input())
    if duoxNum != 0:
        print('请输入多选题分数')
        duoxFs = int(input())
        th.append(['多选', duoxFs])
        timubiaoti = doc.add_paragraph()
        xiaofen = duoxFs / duoxNum
        timubiaotiRun = timubiaoti.add_run(daxie[thSum] + '、多选题' + '（每题' + str(xiaofen) + '分，共计' + str(duoxFs) + '分）')
        timubiaotiRun.font.bold = True

        timubiaoti1 = doc1.add_paragraph()
        timubiaotiRun1 = timubiaoti1.add_run(daxie[thSum] + '、多选题' + '（每题' + str(xiaofen) + '分，共计' + str(duoxFs) + '分）')
        timubiaotiRun1.font.bold = True
        thSum = thSum + 1
        chuti(5, duoxNum, duoxFs)
    print('请输入判断题数量')
    pdNum=int(input())
    if pdNum!=0:
        print('请输入判断题分数')
        pdFs = int(input())
        th.append(['判断',pdFs])
        timubiaoti=doc.add_paragraph()
        xiaofen=pdFs/pdNum
        timubiaotiRun=timubiaoti.add_run(daxie[thSum]+'、判断题'+'（每题'+str(xiaofen)+'分，共计'+str(pdFs)+'分）')
        timubiaotiRun.font.bold=True

        timubiaoti1=doc1.add_paragraph()
        timubiaotiRun1=timubiaoti1.add_run(daxie[thSum]+'、判断题'+'（每题'+str(xiaofen)+'分，共计'+str(pdFs)+'分）')
        timubiaotiRun1.font.bold=True
        thSum = thSum + 1
        chuti(3, pdNum,pdFs)
    print('请输入简答题数量')
    jdNum=int(input())
    if jdNum!=0:
        print('请输入简答题分数')
        jdFs = int(input())
        th.append(['简答',jdFs])
        timubiaoti=doc.add_paragraph()
        xiaofen=jdFs/jdNum
        timubiaotiRun=timubiaoti.add_run(daxie[thSum]+'、填空题'+'（每题'+str(xiaofen)+'分，共计'+str(jdFs)+'分）')
        timubiaotiRun.font.bold=True

        timubiaoti1=doc1.add_paragraph()
        timubiaotiRun1=timubiaoti1.add_run(daxie[thSum]+'、填空题'+'（每题'+str(xiaofen)+'分，共计'+str(jdFs)+'分）')
        timubiaotiRun1.font.bold=True
        thSum = thSum + 1
        chuti(4,jdNum,jdFs)
    for x in th:
        fenshu.add_column(Cm(2)).cells[1].text=x[0]
        fenshu1.add_column(Cm(2)).cells[1].text = x[0]

    fenshu.add_column(Cm(2)).cells[0].text = '总分'
    fenshu.add_column(Cm(2)).cells[0].text = '阅卷人'
    fenshu.cell(0, 1).text = '题号'
    fenshu.cell(0,1).merge(fenshu.cell(0,len(th)))
    fenshu.cell(0,len(th)+1).merge(fenshu.cell(1,len(th)+1))
    fenshu.cell(0, len(th) + 2).merge(fenshu.cell(1, len(th) + 2))
    fenshu.alignment = WD_TABLE_ALIGNMENT.CENTER

    fenshu1.add_column(Cm(2)).cells[0].text = '总分'
    fenshu1.add_column(Cm(2)).cells[0].text = '阅卷人'
    fenshu1.cell(0, 1).text = '题号'
    fenshu1.cell(0,1).merge(fenshu1.cell(0,len(th)))
    fenshu1.cell(0,len(th)+1).merge(fenshu1.cell(1,len(th)+1))
    fenshu1.cell(0, len(th) + 2).merge(fenshu1.cell(1, len(th) + 2))
    fenshu1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in fenshu.rows:
        row.height = Cm(1)
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Cm(30)
    for row in fenshu1.rows:
        row.height = Cm(1)
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Cm(30)
    fenshu.autofit=False
    fenshu1.autofit = False
    #fenshu.autofit = True
    doc.save("答案.docx")
    doc1.save("试卷.docx")